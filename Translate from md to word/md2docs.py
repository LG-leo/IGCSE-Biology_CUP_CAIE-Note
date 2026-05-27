#!/usr/bin/env python3
"""MD → DOCX v8（修复上标内τ错误）"""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SUP_MAP = {
    '0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹',
    '+':'⁺','-':'⁻','=':'⁼','(':'⁽',')':'⁾','n':'ⁿ','i':'ⁱ',
    'a':'ᵃ','b':'ᵇ','c':'ᶜ','d':'ᵈ','e':'ᵉ','f':'ᶠ','g':'ᵍ','h':'ʰ',
    'j':'ʲ','k':'ᵏ','l':'ˡ','m':'ᵐ','o':'ᵒ','p':'ᵖ','r':'ʳ','s':'ˢ',
    't':'ᵗ','u':'ᵘ','v':'ᵛ','w':'ʷ','x':'ˣ','y':'ʸ','z':'ᶻ',
}
SUB_MAP = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅','6':'₆','7':'₇','8':'₈','9':'₉','+':'₊','-':'₋','=':'₌','(':'₍',')':'₎'}

LATEX_REPLACE = [
    (r'\iiint','∭'),(r'\iint','∬'),(r'\oint','∮'),(r'\infty','∞'),(r'\partial','∂'),(r'\nabla','∇'),
    (r'\forall','∀'),(r'\exists','∃'),(r'\varnothing','∅'),(r'\emptyset','∅'),(r'\notin','∉'),
    (r'\subseteq','⊆'),(r'\supseteq','⊇'),(r'\subset','⊂'),(r'\supset','⊃'),
    (r'\rightarrow','→'),(r'\leftarrow','←'),(r'\Rightarrow','⇒'),(r'\Leftarrow','⇐'),
    (r'\Leftrightarrow','⇔'),(r'\leftrightarrow','↔'),(r'\longrightarrow','→'),(r'\longleftarrow','←'),
    (r'\Longrightarrow','⇒'),(r'\Longleftarrow','⇐'),(r'\implies','⇒'),(r'\iff','⇔'),(r'\mapsto','↦'),
    (r'\approx','≈'),(r'\equiv','≡'),(r'\simeq','≃'),(r'\cong','≅'),(r'\sim','∼'),(r'\propto','∝'),
    (r'\times','×'),(r'\div','÷'),(r'\pm','±'),(r'\mp','∓'),(r'\cdot','·'),(r'\neq','≠'),(r'\leq','≤'),(r'\geq','≥'),
    (r'\ll','≪'),(r'\gg','≫'),(r'\oplus','⊕'),(r'\otimes','⊗'),(r'\ominus','⊖'),(r'\odot','⊙'),(r'\oslash','⊘'),
    (r'\land','∧'),(r'\lor','∨'),(r'\neg','¬'),(r'\in','∈'),(r'\ni','∋'),(r'\cup','∪'),(r'\cap','∩'),
    (r'\setminus','∖'),(r'\angle','∠'),(r'\perp','⊥'),(r'\parallel','∥'),(r'\triangle','△'),(r'\square','□'),
    (r'\hbar','ℏ'),(r'\ell','ℓ'),(r'\imath','ı'),(r'\jmath','ȷ'),(r'\Re','ℜ'),(r'\Im','ℑ'),(r'\aleph','ℵ'),(r'\wp','℘'),
    (r'\prime','′'),(r'\to','→'),(r'\top','⊤'),(r'\bot','⊥'),
    (r'\int','∫'),(r'\sum','∑'),(r'\prod','∏'),(r'\coprod','∐'),
    (r'\sinh','sinh'),(r'\cosh','cosh'),(r'\tanh','tanh'),
    (r'\arcsin','arcsin'),(r'\arccos','arccos'),(r'\arctan','arctan'),
    (r'\sin','sin'),(r'\cos','cos'),(r'\tan','tan'),(r'\cot','cot'),(r'\sec','sec'),(r'\csc','csc'),
    (r'\log','log'),(r'\ln','ln'),(r'\lg','lg'),(r'\lim','lim'),(r'\max','max'),(r'\min','min'),
    (r'\sup','sup'),(r'\inf','inf'),(r'\det','det'),(r'\deg','deg'),(r'\arg','arg'),
    (r'\ldots','…'),(r'\cdots','…'),(r'\vdots','⋮'),(r'\ddots','⋱'),
    (r'\alpha','α'),(r'\beta','β'),(r'\gamma','γ'),(r'\delta','δ'),(r'\epsilon','ε'),(r'\varepsilon','ε'),
    (r'\zeta','ζ'),(r'\eta','η'),(r'\theta','θ'),(r'\vartheta','θ'),
    (r'\iota','ι'),(r'\kappa','κ'),(r'\lambda','λ'),(r'\mu','μ'),(r'\nu','ν'),(r'\xi','ξ'),
    (r'\pi','π'),(r'\varpi','ϖ'),(r'\rho','ρ'),(r'\varrho','ρ'),(r'\sigma','σ'),(r'\varsigma','ς'),
    (r'\tau','τ'),(r'\upsilon','υ'),(r'\phi','φ'),(r'\varphi','φ'),(r'\chi','χ'),(r'\psi','ψ'),(r'\omega','ω'),
    (r'\Gamma','Γ'),(r'\Delta','Δ'),(r'\Theta','Θ'),(r'\Lambda','Λ'),(r'\Xi','Ξ'),(r'\Pi','Π'),
    (r'\Sigma','Σ'),(r'\Phi','Φ'),(r'\Psi','Ψ'),(r'\Omega','Ω'),
    (r'\mathbb{N}','ℕ'),(r'\mathbb{Z}','ℤ'),(r'\mathbb{Q}','ℚ'),(r'\mathbb{R}','ℝ'),(r'\mathbb{C}','ℂ'),
]

def extract_braces(s, pos):
    if pos >= len(s) or s[pos] != '{': return None, pos
    depth, i = 0, pos
    while i < len(s):
        if s[i] == '{': depth += 1
        elif s[i] == '}':
            depth -= 1
            if depth == 0: return s[pos+1:i], i+1
        i += 1
    return None, pos

def replace_all_cmds(s):
    for cmd, uni in sorted(LATEX_REPLACE, key=lambda x: -len(x[0])):
        s = s.replace(cmd, uni)
    return s

def latex_to_text(s):
    if not s or not s.strip(): return ''
    s = s.strip()
    s = re.sub(r'\\\(', '', s); s = re.sub(r'\\\)', '', s)
    s = re.sub(r'\\\[', '', s); s = re.sub(r'\\\]', '', s)
    for env in ['aligned','align','split','gathered','cases','matrix','pmatrix','bmatrix','vmatrix','array','eqnarray']:
        s = re.sub(r'\\begin\{'+env+r'\}[t\*]?\s*', '', s)
        s = re.sub(r'\\end\{'+env+r'\}[t\*]?\s*', '', s)
    s = re.sub(r'\\begin\{[^}]*\}', '', s); s = re.sub(r'\\end\{[^}]*\}', '', s)
    s = s.replace('\\\\', '\n').replace('\\cr', '\n')
    s = s.replace('&', '')
    
    while '\\boxed' in s:
        idx = s.find('\\boxed'); rest = s[idx+6:]
        if rest and rest[0] == '{':
            cnt, np = extract_braces(rest, 0)
            if cnt is not None: s = s[:idx] + cnt + rest[np:]; continue
        break
    
    while '\\binom' in s:
        idx = s.find('\\binom'); rest = s[idx+6:]
        if rest and rest[0] == '{':
            top, p1 = extract_braces(rest, 0)
            if top and p1 < len(rest) and rest[p1] == '{':
                bot, p2 = extract_braces(rest, p1)
                if bot: s = s[:idx] + f'C({latex_to_text(top)},{latex_to_text(bot)})' + rest[p2:]; continue
        break
    
    s = s.replace('\\dfrac','\\frac').replace('\\tfrac','\\frac').replace('\\cfrac','\\frac')
    
    while '\\frac' in s:
        idx = s.find('\\frac'); rest = s[idx+5:]
        if rest and rest[0] == '{':
            num, p1 = extract_braces(rest, 0)
            if num and p1 < len(rest) and rest[p1] == '{':
                den, p2 = extract_braces(rest, p1)
                if den: s = s[:idx] + f'({latex_to_text(num)}/{latex_to_text(den)})' + rest[p2:]; continue
        break
    
    while '\\sqrt' in s:
        idx = s.find('\\sqrt'); rest = s[idx+5:]; n = None; pos = 0
        if rest and rest[0] == '[':
            eb = rest.find(']')
            if eb > 0: n = rest[1:eb]; pos = eb + 1
        if pos < len(rest) and rest[pos] == '{':
            body, p2 = extract_braces(rest, pos)
            if body: s = s[:idx] + (f'{n}√({latex_to_text(body)})' if n else f'√({latex_to_text(body)})') + rest[p2:]; continue
        break
    
    for cmd in ['mathbf','mathrm','mathit','mathsf','mathtt','mathbb','mathcal','mathscr','mathfrak','bm','textbf','textit','text','textrm','emph']:
        while True:
            idx = s.find(f'\\{cmd}')
            if idx == -1: break
            rest = s[idx+len(cmd)+1:]
            if rest and rest[0] == '{':
                cnt, np = extract_braces(rest, 0)
                if cnt is not None: s = s[:idx] + cnt + rest[np:]; continue
            break
    
    for dec in ['widehat','widetilde','overleftrightarrow','overrightarrow','overleftarrow',
                'hat','vec','dot','ddot','tilde','bar','overline','underline']:
        while True:
            idx = s.find(f'\\{dec}')
            if idx == -1: break
            rest = s[idx+len(dec)+1:]
            if rest and rest[0] == '{':
                cnt, np = extract_braces(rest, 0)
                if cnt is not None: s = s[:idx] + cnt + rest[np:]; continue
            if rest and re.match(r'[a-zA-Z]', rest[0]):
                s = s[:idx] + rest[0] + rest[1:]; continue
            break
    
    for cmd in ['left','right','bigl','bigr','big','Bigl','Bigr','Big','biggl','biggr','bigg','Biggl','Biggr','Bigg']:
        s = re.sub(rf'\\{cmd}\s*', '', s)
    for cmd in ['displaystyle','textstyle','scriptstyle','scriptscriptstyle','limits','nolimits','vcenter','hbox','mbox']:
        s = re.sub(rf'\\{cmd}\s*', '', s)
    for sp in [r'\\,\s*',r'\\!\s*',r'\\;\s*',r'\\:\s*']: s = re.sub(sp, '', s)
    s = re.sub(r'\\quad\s*', '  ', s); s = re.sub(r'\\qquad\s*', '    ', s)
    s = re.sub(r'\\hspace\{[^}]*\}', '', s); s = re.sub(r'\\vspace\{[^}]*\}', '', s)
    s = re.sub(r'\\color\{[^}]*\}', '', s); s = re.sub(r'\\textcolor\{[^}]*\}\{([^}]*)\}', r'\1', s)
    s = re.sub(r'\\overrightarrow\{([^}]*)\}', r'\1→', s)
    s = re.sub(r'\\overleftarrow\{([^}]*)\}', r'←\1', s)
    s = re.sub(r'\\underrightarrow\{([^}]*)\}', r'\1→', s)
    s = re.sub(r'\\underbrace\{([^}]*)\}\{([^}]*)\}', r'\1', s)
    s = re.sub(r'\\overbrace\{([^}]*)\}\{([^}]*)\}', r'\1', s)
    
    # ⭐ 关键修复：先替换所有 LaTeX 命令（将 \tau 变成 τ），再处理上下标
    s = replace_all_cmds(s)
    
    # 然后处理上下标（此时 \tau 已变成 τ，不会被错误转换）
    s = re.sub(r'\^\{(.*?)\}', lambda m: ''.join(SUP_MAP.get(c,c) for c in m.group(1)), s)
    s = re.sub(r'_\{(.*?)\}', lambda m: ''.join(SUB_MAP.get(c,c) for c in m.group(1)), s)
    s = re.sub(r'\^([a-zA-Z0-9])', lambda m: SUP_MAP.get(m.group(1), f'^{m.group(1)}'), s)
    s = re.sub(r'_([a-zA-Z0-9])', lambda m: SUB_MAP.get(m.group(1), f'_{m.group(1)}'), s)
    
    # 残余命令
    s = re.sub(r'\\([a-zA-Z]+)\{([^}]*)\}', r'\2', s)
    s = re.sub(r'\\([a-zA-Z]+)', '', s)
    s = re.sub(r'\{(\w)\}', r'\1', s)
    s = s.replace('{','').replace('}','').replace('~',' ')
    
    lines = s.split('\n')
    lines = [re.sub(r'\s+', ' ', l).strip() for l in lines]
    lines = [l for l in lines if l]
    return '\n'.join(lines)


def split_inline(text):
    result = []; pos, n = 0, len(text)
    while pos < n:
        d1 = text.find('$', pos); d2 = text.find(r'\(', pos)
        if d1 == -1 and d2 == -1:
            result.append(('text', text[pos:])); break
        elif d1 == -1: d, st = d2, 'paren'
        elif d2 == -1: d, st = d1, 'dollar'
        else: d, st = (d1,'dollar') if d1<=d2 else (d2,'paren')
        if st=='dollar' and d>0 and text[d-1]=='\\':
            result.append(('text', text[pos:d-1]+'$')); pos=d+1; continue
        if d > pos: result.append(('text', text[pos:d]))
        if st == 'dollar':
            if d+1<n and text[d+1]=='$': pos=d+2; continue
            e = text.find('$',d+1)
            if e == -1: pos=n; break
            content = text[d+1:e]
            if content.strip() and '\n' not in content:
                result.append(('math', latex_to_text(content.strip())))
            pos = e+1
        else:
            e = text.find(r'\)', d+2)
            if e == -1: pos=n; break
            content = text[d+2:e]
            if content.strip() and '\n' not in content:
                result.append(('math', latex_to_text(content.strip())))
            pos = e+2
    return result


def add_formatted_text(paragraph, text):
    pos, n = 0, len(text)
    while pos < n:
        db = text.find('**', pos); sb = text.find('*', pos)
        candidates = []
        if db != -1: candidates.append((db, 'db'))
        if sb != -1: candidates.append((sb, 'sb'))
        if not candidates: paragraph.add_run(text[pos:]); break
        candidates.sort()
        first_pos, first_type = candidates[0]
        if first_type == 'db':
            if first_pos > pos: paragraph.add_run(text[pos:first_pos])
            end = text.find('**', first_pos+2)
            if end == -1: paragraph.add_run(text[first_pos:]); break
            run = paragraph.add_run(text[first_pos+2:end]); run.bold = True; pos = end+2
        else:
            if sb > pos: paragraph.add_run(text[pos:sb])
            end = text.find('*', sb+1)
            if end == -1: paragraph.add_run(text[sb:]); break
            if end > sb+1 and text[end-1] == '*':
                paragraph.add_run(text[sb:end+1]); pos = end+1
            else:
                run = paragraph.add_run(text[sb+1:end]); run.italic = True; pos = end+1

def is_table_row(line):
    s = line.strip(); return s.startswith('|') and s.count('|')>=2
def is_table_separator(line):
    return bool(re.match(r'^\|[\s:-]+\|', line.strip()))
def parse_table_row(line):
    s = line.strip()
    if s.startswith('|'): s = s[1:]
    if s.endswith('|'): s = s[:-1]
    return [c.strip() for c in s.split('|')]
def get_table_alignments(sep_row):
    cells = parse_table_row(sep_row); aligns = []
    for cell in cells:
        if cell.startswith(':') and cell.endswith(':'): aligns.append('center')
        elif cell.endswith(':'): aligns.append('right')
        elif cell.startswith(':'): aligns.append('left')
        else: aligns.append('left')
    return aligns

def parse_md(text):
    paras = []; lines = text.split('\n'); i, n = 0, len(lines)
    while i < n:
        s = lines[i].strip()
        if s.startswith('```'):
            code = [lines[i]]; i+=1
            while i<n and not lines[i].strip().startswith('```'): code.append(lines[i]); i+=1
            if i<n: code.append(lines[i]); i+=1
            paras.append(('code',0,'\n'.join(code))); continue
        if is_table_row(s) and i+1<n and is_table_separator(lines[i+1]):
            header = parse_table_row(s); i+=1
            aligns = get_table_alignments(lines[i]); i+=1
            data_rows = []
            while i<n and is_table_row(lines[i]):
                cells = parse_table_row(lines[i])
                processed = [''.join(s[1] for s in split_inline(c)) for c in cells]
                data_rows.append(processed); i+=1
            hp = [''.join(s[1] for s in split_inline(c)) for c in header]
            paras.append(('table',0,(hp, data_rows, aligns))); continue
        if s.startswith('$$') and s.endswith('$$') and len(s)>4:
            r = latex_to_text(s[2:-2])
            if r: paras.append(('display',0,r)); i+=1; continue
        if s.startswith('$$'):
            tex=[]; i+=1
            while i<n and '$$' not in lines[i]: tex.append(lines[i]); i+=1
            if i<n:
                extra = lines[i].replace('$$','').strip()
                if extra: tex.append(extra); i+=1
            t='\n'.join(tex).strip()
            if t:
                r=latex_to_text(t)
                if r: paras.append(('display',0,r))
            continue
        if s.startswith(r'\['):
            tex=[s[2:].strip()]; i+=1
            while i<n and r'\]' not in lines[i]: tex.append(lines[i]); i+=1
            if i<n:
                extra=lines[i].replace(r'\]','').strip()
                if extra: tex[-1]=extra; i+=1
            t='\n'.join(tex).strip()
            if t:
                r=latex_to_text(t)
                if r: paras.append(('display',0,r))
            continue
        h=re.match(r'^(#{1,6})\s+(.*)',lines[i])
        if h:
            segs = split_inline(h.group(2)); text = ''.join(s[1] for s in segs)
            if text: paras.append(('heading',len(h.group(1)),text)); i+=1; continue
        if not s: i+=1; continue
        pl=[lines[i]]; i+=1
        while i<n and lines[i].strip():
            ns=lines[i].strip()
            if re.match(r'^#{1,6}\s',lines[i]) or ns.startswith('$$') or ns.startswith('```') or ns.startswith(r'\[') or (is_table_row(ns) and i+1<n and is_table_separator(lines[i+1])):
                break
            pl.append(lines[i]); i+=1
        pt=re.sub(r' +',' ',' '.join(l.strip() for l in pl if l.strip()))
        segs = split_inline(pt); text = ''.join(s[1] for s in segs)
        if text: paras.append(('para',0,text))
    return paras

def create_docx(paras, path):
    doc = Document()
    for kind, level, content in paras:
        if kind == 'display':
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(p, content)
        elif kind == 'heading':
            p=doc.add_heading(level=min(level,4))
            add_formatted_text(p, content)
        elif kind == 'para':
            p=doc.add_paragraph()
            add_formatted_text(p, content)
        elif kind == 'code':
            p=doc.add_paragraph(); r=p.add_run(content); r.font.name='Courier New'; r.font.size=Pt(9)
        elif kind == 'table':
            header,data_rows,aligns=content; nc=len(header)
            tbl=doc.add_table(rows=1+len(data_rows),cols=nc); tbl.style='Table Grid'
            for j,ct in enumerate(header):
                c=tbl.rows[0].cells[j]; add_formatted_text(c.paragraphs[0], ct)
                for p2 in c.paragraphs:
                    for r2 in p2.runs: r2.bold=True
                if j<len(aligns):
                    m={'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER,'right':WD_ALIGN_PARAGRAPH.RIGHT}
                    c.paragraphs[0].alignment=m.get(aligns[j],WD_ALIGN_PARAGRAPH.LEFT)
            for i,row in enumerate(data_rows):
                for j,ct in enumerate(row):
                    if j>=nc: break
                    c=tbl.rows[i+1].cells[j]; add_formatted_text(c.paragraphs[0], ct)
                    if j<len(aligns):
                        m={'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER,'right':WD_ALIGN_PARAGRAPH.RIGHT}
                        c.paragraphs[0].alignment=m.get(aligns[j],WD_ALIGN_PARAGRAPH.LEFT)
            doc.add_paragraph()
    doc.save(path)
    print(f'✅ DOCX → {path}')

if __name__ == '__main__':
    if len(sys.argv)<2: print('用法: python3 md2docs.py input.md'); sys.exit(1)
    inp=Path(sys.argv[1])
    if not inp.exists(): print(f'❌ 文件不存在: {sys.argv[1]}'); sys.exit(1)
    paras=parse_md(inp.read_text(encoding='utf-8'))
    create_docx(paras,inp.stem+'.docx')
